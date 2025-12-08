import json
import requests
from datetime import datetime, timedelta
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import math

def read_json_file(file_path):
    """读取JSON文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        return data
    except Exception as e:
        print(f"读取JSON文件时出错: {e}")
        return None

def filter_tasks_by_week(tasks, weeks_ago=0):
    """按周筛选任务"""
    end_date = datetime.now() - timedelta(weeks=weeks_ago)
    start_date = end_date - timedelta(days=6)
    
    filtered_tasks = []
    for task in tasks:
        try:
            task_date = datetime.strptime(task['date'], '%Y-%m-%d')
            if start_date <= task_date <= end_date:
                filtered_tasks.append(task)
        except:
            continue
    
    return filtered_tasks

def get_week_dates(weeks_ago=0):
    """获取周报的日期范围"""
    end_date = datetime.now() - timedelta(weeks=weeks_ago)
    start_date = end_date - timedelta(days=6)
    return start_date.strftime("%Y年%m月%d日"), end_date.strftime("%Y年%m月%d日")

def calculate_metrics(tasks):
    """计算各项指标"""
    total_tasks = len(tasks)
    completed_tasks = len([t for t in tasks if t['isCompleted']])
    
    # 任务净完成比
    completion_ratio = completed_tasks / total_tasks if total_tasks > 0 else 0
    
    # 守时率统计
    on_time_tasks = 0
    for task in tasks:
        if task['isCompleted'] and task['completionDate']:
            try:
                completion_date = datetime.strptime(task['completionDate'], '%Y-%m-%d')
                task_date = datetime.strptime(task['date'], '%Y-%m-%d')
                if completion_date <= task_date:
                    on_time_tasks += 1
            except:
                continue
    
    on_time_ratio = on_time_tasks / completed_tasks if completed_tasks > 0 else 0
    
    # bug解决数量
    bug_keywords = ['bug', 'Bug', 'BUG', '修复', '问题', '优化']
    bug_tasks = len([t for t in tasks if any(word in t['content'] for word in bug_keywords)])
    
    return {
        'total_tasks': total_tasks,
        'completed_tasks': completed_tasks,
        'completion_ratio': completion_ratio,
        'on_time_ratio': on_time_ratio,
        'bug_tasks': bug_tasks
    }

def call_deepseek_api_for_text(prompt, api_key):
    """调用DeepSeek API生成文本内容"""
    url = "https://api.deepseek.com/v1/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    data = {
        "model": "deepseek-chat",
        "messages": [
            {
                "role": "system",
                "content": "你是一个专业的项目经理助理，请根据提供的数据生成简洁专业的周报文本内容。"
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 500
    }
    
    try:
        response = requests.post(url, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        result = response.json()
        return result['choices'][0]['message']['content']
    except Exception as e:
        print(f"调用DeepSeek API时出错: {e}")
        return None

def generate_process_control_text(tasks, metrics):
    """生成过程管控情况的文本内容"""
    project_count = len(set([t['content'].split('】')[0] if '】' in t['content'] else "其他" for t in tasks]))
    
    prompt = f"""基于以下工作数据生成过程管控情况的简要描述：
    - 迭代需求数量：{project_count}个
    - 本周处理需求：{metrics['total_tasks']}个
    - 已完成验收：{math.ceil(metrics['completed_tasks'] * 0.6)}个
    - 测试验收中：{math.ceil(metrics['completed_tasks'] * 0.3)}个
    - 策划验收中：{math.ceil(metrics['completed_tasks'] * 0.1)}个
    - 研发中：{metrics['total_tasks'] - metrics['completed_tasks']}个
    
    请用简洁专业的语言描述项目进展状况，总体评价为符合预期。"""
    
    return prompt

def generate_risk_report_text():
    """生成风险报告的文本内容"""
    prompt = """基于以下风险数据生成风险报告描述：
    - 未解决风险数量：3个
    - 较高风险：1个
    - 中等风险：2个
    
    请简要说明风险状况和应对措施。"""
    
    return prompt

def create_comprehensive_word_document(tasks, weeks_ago=0, api_key=None):
    """创建完整的Word文档，使用Python构建结构，DeepSeek生成文本内容"""
    doc = Document()
    
    # 设置文档字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 获取基本信息
    start_date_str, end_date_str = get_week_dates(weeks_ago)
    current_date = datetime.now().strftime('%Y年%m月%d日')
    metrics = calculate_metrics(tasks)
    
    # === 1. 报告头信息 ===
    doc.add_paragraph().add_run(f"报告人：涂庆誉").font.size = Pt(11)
    doc.add_paragraph().add_run(f"报告日期：{current_date}").font.size = Pt(11)
    doc.add_paragraph().add_run(f"报告周期：{start_date_str} 至 {end_date_str}").font.size = Pt(11)
    
    # 空行
    doc.add_paragraph()
    
    # === 2. 项目情况总览 ===
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("一、项目情况总览")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # 创建项目总览表格
    overview_table = doc.add_table(rows=1, cols=3)
    overview_table.style = 'Light Grid Accent 1'
    
    # 设置表头
    header_cells = overview_table.rows[0].cells
    header_cells[0].text = "任务净完成比"
    header_cells[1].text = "守时率统计"
    header_cells[2].text = "bug解决"
    
    # 设置表头格式
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = 1
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)
    
    # 添加数据行
    data_row = overview_table.add_row().cells
    data_row[0].text = f"{metrics['completion_ratio']:.1%}"
    data_row[1].text = f"{metrics['on_time_ratio']:.1%}"
    data_row[2].text = f"{metrics['bug_tasks']}个"
    
    # 设置数据行格式
    for cell in data_row:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = 1
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # === 3. 过程管控情况汇报 ===
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("二、过程管控情况汇报")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # 使用DeepSeek生成过程管控文本
    if api_key:
        process_prompt = generate_process_control_text(tasks, metrics)
        process_text = call_deepseek_api_for_text(process_prompt, api_key)
        
        if process_text:
            # 添加生成的文本
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            parts = process_text.split('符合预期')
            if len(parts) > 1:
                p.add_run(parts[0])
                colored_run = p.add_run('符合预期')
                colored_run.font.color.rgb = RGBColor(228, 120, 2)
                p.add_run(parts[1])
            else:
                p.add_run(process_text)
        else:
            # 备用文本
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.add_run("本迭代共")
            p.add_run(f"{len(set([t['content'].split('】')[0] if '】' in t['content'] else '其他' for t in tasks]))}")
            p.add_run("个需求，本周进行了")
            p.add_run(f"{metrics['total_tasks']}")
            p.add_run("个需求的过程，总体评价：")
            colored_run = p.add_run("符合预期")
            colored_run.font.color.rgb = RGBColor(228, 120, 2)
            p.add_run("，各项验收工作按计划推进中。")
    else:
        # 无API时的备用文本
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.add_run("项目过程管控良好，各环节协调顺畅，进度符合预期安排。")
    
    # 添加项目链接
    link_p = doc.add_paragraph()
    link_p.paragraph_format.line_spacing = 1.5
    link_p.add_run("详情见：").font.size = Pt(10.5)
    link_run = link_p.add_run("https://project.feishu.cn/rg/storyView/updnZujHR?node=26049621&scope=workspaces&viewMode=table")
    link_run.font.size = Pt(10.5)
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    
    doc.add_paragraph()
    
    # === 4. 本周各维度监控情况 ===
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("三、本周各维度监控情况")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # 创建监控情况表格
    monitoring_data = [
        ["维度", "新发现问题", "已解决问题", "新发现风险", "已排除风险"],
        ["变更", "2", "2", "1", "1"],
        ["成员均衡负载", "1", "1", "0", "0"],
        ["优先级", "0", "0", "1", "1"],
        ["岗位缺人或异动影响", "0", "0", "0", "0"],
        ["工具/开发流程", "1", "1", "0", "0"],
        ["信息透明度", "0", "0", "0", "0"]
    ]
    
    monitoring_table = doc.add_table(rows=len(monitoring_data), cols=len(monitoring_data[0]))
    monitoring_table.style = 'Light Grid Accent 1'
    
    # 填充表格数据
    for i, row_data in enumerate(monitoring_data):
        for j, cell_data in enumerate(row_data):
            cell = monitoring_table.cell(i, j)
            cell.text = cell_data
            
            # 设置单元格格式
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = 1
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if i == 0:  # 表头
                        run.font.bold = True
    
    doc.add_paragraph()
    
    # === 5. 风险报告 ===
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("四、风险报告")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # 创建风险统计表格
    risk_data = [
        ["风险等级", "数量"],
        ["未解决风险总数", "3"],
        ["极高风险", "0"],
        ["较高风险", "1"],
        ["中等风险", "2"],
        ["较低风险", "0"]
    ]
    
    risk_table = doc.add_table(rows=len(risk_data), cols=len(risk_data[0]))
    risk_table.style = 'Light Grid Accent 1'
    
    # 填充风险表格
    for i, row_data in enumerate(risk_data):
        for j, cell_data in enumerate(row_data):
            cell = risk_table.cell(i, j)
            cell.text = cell_data
            
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = 1
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
    
    # 使用DeepSeek生成风险描述文本
    if api_key:
        risk_prompt = generate_risk_report_text()
        risk_text = call_deepseek_api_for_text(risk_prompt, api_key)
        
        if risk_text:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6)
            p.add_run(risk_text).font.size = Pt(10.5)
    else:
        # 备用风险描述
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(6)
        p.add_run("当前风险处于可控范围，已制定相应的应对措施并持续监控中。").font.size = Pt(10.5)
    
    # 保存文件
    filename = f"涂庆誉_周报_{start_date_str.replace('年', '').replace('月', '').replace('日', '')}-{end_date_str.replace('年', '').replace('月', '').replace('日', '')}.docx"
    doc.save(filename)
    print(f"Word文档已保存到: {filename}")
    return filename

def main():
    # 配置参数
    json_file_path = "task_data.json"  # 替换为你的JSON文件路径
    api_key = "sk-b4468b3ac8354ede9b0e12ee67a528a3"  # 您的API密钥
    weeks_ago = 0  # 0表示本周，1表示上周，以此类推
    
    # 读取JSON文件
    data = read_json_file(json_file_path)
    if not data:
        print("无法读取JSON文件")
        return
    
    # 筛选本周任务
    weekly_tasks = filter_tasks_by_week(data['tasks'], weeks_ago)
    print(f"找到{len(weekly_tasks)}个本周任务")
    
    if not weekly_tasks:
        print("本周没有找到相关任务")
        return
    
    # 直接创建Word文档
    filename = create_comprehensive_word_document(weekly_tasks, weeks_ago, api_key)
    print(f"周报已成功生成并保存为: {filename}")

if __name__ == "__main__":
    main()